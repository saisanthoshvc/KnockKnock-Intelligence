import os
import json
import time
import io
from datetime import datetime, timedelta

import requests
from flask import Flask, render_template, request, jsonify, redirect, url_for, send_file
from werkzeug.utils import secure_filename

# =========================================================
# Flask app
# =========================================================
app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'your-secret-key-here')

# ---------- Paths (use /tmp on Posit Connect) ----------
UPLOAD_FOLDER = os.environ.get('UPLOAD_DIR', '/tmp/knockknock_uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
STATS_FILE = os.environ.get('STATS_FILE', '/tmp/knockknock_stats.json')

# ---------- Allowed uploads ----------
ALLOWED_EXTENSIONS = {'docx', 'pdf', 'xlsx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB

# =========================================================
# Helpers
# =========================================================
def env_int(name, default):
    try:
        return int(os.environ.get(name, default))
    except Exception:
        return default

def log_event(event, **kv):
    """Prints structured JSON logs (friendly to Posit Connect)."""
    try:
        print(json.dumps({"event": event, **kv}))
    except Exception:
        print(f"[{event}] {kv}")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# =========================================================
# EXTRACTOR (RFP) API CONFIG (unchanged from your setup)
# =========================================================
DEFAULT_API_CONFIG = {
    'base_url': os.environ.get(
        'RFP_API_URL',
        # Keep your previous default so existing extractor continues to work
        'https://connect.affiniusaiplatform.com/content/000d3572-937a-4d5c-ab7e-7ec80d80c4ce/'
    ),
    'api_key': os.environ.get('RFP_API_KEY', ''),
    'header_name': os.environ.get('RFP_API_KEY_HEADER_NAME', 'Authorization'),
    'auth_scheme': os.environ.get('RFP_API_AUTH_SCHEME', 'Key'),  # Key | Bearer | Raw
    'poll_interval': env_int('POLL_INTERVAL_SECONDS', 2),
    'poll_timeout': env_int('POLL_TIMEOUT_SECONDS', 240),
    'request_timeout': env_int('REQUEST_TIMEOUT_SECONDS', 60),
}
api_config = DEFAULT_API_CONFIG.copy()

def get_auth_header():
    """
    Build extractor API auth header according to env:
      - Authorization: Key <token>
      - Authorization: Bearer <token>
      - <custom-name>: <token>
    """
    hdr = api_config['header_name']
    scheme = api_config['auth_scheme'].lower()
    token = api_config['api_key']

    if not token:
        return {}

    if hdr.lower() == 'authorization':
        if scheme == 'key':
            return {'Authorization': f'Key {token}'}
        elif scheme == 'bearer':
            return {'Authorization': f'Bearer {token}'}
        elif scheme == 'raw':
            return {'Authorization': token}
        else:
            return {'Authorization': token}
    else:
        if scheme == 'raw':
            return {hdr: token}
        elif scheme in ('key', 'bearer'):
            return {hdr: f'{scheme.title()} {token}'}
        else:
            return {hdr: token}

def make_api_request(method, endpoint, **kwargs):
    url = f"{api_config['base_url'].rstrip('/')}/{endpoint.strip('/')}"
    headers = get_auth_header()
    headers.update(kwargs.pop('headers', {}))

    log_event("API_REQUEST_OUT", method=method.upper(), url=url, header_name=api_config['header_name'])

    try:
        if method.upper() == 'GET':
            resp = requests.get(url, headers=headers, timeout=api_config['request_timeout'])
        elif method.upper() == 'POST':
            resp = requests.post(url, headers=headers, timeout=api_config['request_timeout'], **kwargs)
        elif method.upper() == 'DELETE':
            resp = requests.delete(url, headers=headers, timeout=api_config['request_timeout'])
        else:
            return None

        # Detect redirect to login page (auth failure)
        ct = resp.headers.get('Content-Type', '')
        if ('text/html' in ct) and ('__login__' in resp.text.lower() or '<title>sign in' in resp.text.lower()):
            class Dummy:
                status_code = 401
                text = "Not authorized: redirected to login page. Check RFP_API_* env vars."
                def json(self): return {"error": self.text}
            log_event("API_REQUEST_DONE", url=url, status_code=401)
            return Dummy()

        log_event("API_REQUEST_DONE", url=url, status_code=resp.status_code)
        return resp
    except requests.exceptions.RequestException as e:
        log_event("API_REQUEST_ERROR", url=url, error=str(e))
        return None

# =========================================================
# KB API CONFIG (NEW) — separate from extractor
# =========================================================
CONNECT_BASE = os.environ.get('CONNECT_BASE', 'https://connect.affiniusaiplatform.com').rstrip('/')

KK_KB_CONTENT_GUID = os.environ.get('KK_KB_CONTENT_GUID', '').strip()
KK_KB_BASE = (os.environ.get('KK_KB_BASE') or
              (f"{CONNECT_BASE}/content/{KK_KB_CONTENT_GUID}" if KK_KB_CONTENT_GUID else '')).rstrip('/')

# KB Auth (falls back to extractor creds if not provided)
KK_KB_API_KEY = os.environ.get('KK_KB_API_KEY') or os.environ.get('RFP_API_KEY', '')
KK_KB_KEY_HEADER_NAME = os.environ.get('KK_KB_KEY_HEADER_NAME', os.environ.get('RFP_API_KEY_HEADER_NAME', 'Authorization'))
KK_KB_AUTH_SCHEME = os.environ.get('KK_KB_AUTH_SCHEME', os.environ.get('RFP_API_AUTH_SCHEME', 'Key'))
KK_KB_TIMEOUT = env_int('KK_KB_TIMEOUT_SECONDS', 60)

# Optional explicit route override; otherwise we discover from openapi.json
KK_KB_ROUTE = os.environ.get('KK_KB_ROUTE', '').strip()

_cached_kb_route = None  # cache discovered KB route

def get_kb_auth_header():
    """Auth header for the KB content; uses KB vars, falls back to extractor if unset."""
    hdr = (KK_KB_KEY_HEADER_NAME or 'Authorization').strip()
    scheme = (KK_KB_AUTH_SCHEME or 'Key').lower()
    token = KK_KB_API_KEY

    if not token:
        return {}

    if hdr.lower() == 'authorization':
        if scheme == 'key':
            return {'Authorization': f'Key {token}'}
        if scheme == 'bearer':
            return {'Authorization': f'Bearer {token}'}
        if scheme == 'raw':
            return {'Authorization': token}
        return {'Authorization': token}
    else:
        if scheme == 'raw':
            return {hdr: token}
        if scheme in ('key', 'bearer'):
            return {hdr: f'{scheme.title()} {token}'}
        return {hdr: token}

def kb_health():
    if not KK_KB_BASE:
        log_event("KB_HEALTH_SKIP", reason="KB base not configured")
        return 503, "KB base not configured"
    try:
        url = f"{KK_KB_BASE}/health/ready"
        log_event("KB_HEALTH_OUT", url=url)
        r = requests.get(url, headers=get_kb_auth_header(), timeout=15)
        log_event("KB_HEALTH_BACK", status=r.status_code, ct=r.headers.get('Content-Type',''))
        return r.status_code, r.text
    except Exception as e:
        log_event("KB_HEALTH_ERR", error=str(e))
        return 503, str(e)

def discover_kb_route():
    """Return a KB route string like '/kb/query' by inspecting openapi.json, or env override."""
    global _cached_kb_route
    if _cached_kb_route:
        return _cached_kb_route

    if KK_KB_ROUTE:
        route = KK_KB_ROUTE if KK_KB_ROUTE.startswith('/') else f"/{KK_KB_ROUTE}"
        _cached_kb_route = route
        log_event("KB_ROUTE_FROM_ENV", route=route)
        return route

    if not KK_KB_BASE:
        log_event("KB_ROUTE_SKIP", reason="KB base not configured")
        return None

    spec_url = f"{KK_KB_BASE}/openapi.json"
    try:
        log_event("KB_OPENAPI_OUT", url=spec_url)
        spec = requests.get(spec_url, headers=get_kb_auth_header(), timeout=15)
        log_event("KB_OPENAPI_BACK", status=spec.status_code, ct=spec.headers.get('Content-Type',''))

        if spec.status_code != 200:
            return None

        js = spec.json()
        paths = js.get("paths", {})
        # Common candidates — adjust if your spec differs
        candidates = ["/kb/query", "/kb/query-one", "/kb/query_batch", "/api/kb/query", "/query"]

        for p in candidates:
            if p in paths:
                _cached_kb_route = p
                os.environ["KK_KB_ROUTE"] = p  # helpful in later logs
                log_event("KB_ROUTE_DISCOVERED", route=p)
                return p

        # Fallback: first POST path we find
        for p, meta in paths.items():
            methods = {k.lower() for k in meta.keys()}
            if "post" in methods:
                _cached_kb_route = p
                os.environ["KK_KB_ROUTE"] = p
                log_event("KB_ROUTE_FALLBACK", route=p)
                return p

    except Exception as e:
        log_event("KB_OPENAPI_ERR", error=str(e))

    log_event("KB_ROUTE_NONE")
    return None

# =========================================================
# In-memory state
# =========================================================
job_data = {}                # job_id -> {'status', 'questions', 'timestamp', ...}
job_edits = {}               # job_id -> [edited question dicts]
job_responses = {}           # job_id -> list (optional cache)
job_responses_edits = {}     # job_id -> [ {qid, text, response, status, sources} ]

# =========================================================
# Stats helpers
# =========================================================
def load_stats():
    try:
        if os.path.exists(STATS_FILE):
            with open(STATS_FILE, 'r') as f:
                stats = json.load(f)
            docs = stats.get("documents_processed", 0)
            stats["avg_processing_time"] = round(stats.get("total_processing_time", 0) / docs, 1) if docs > 0 else 0.0
            return stats
    except Exception:
        pass
    return {
        "documents_processed": 0,
        "questions_extracted": 0,
        "total_processing_time": 0,
        "avg_processing_time": 0.0,
        "accuracy_rate": 0.0,
        "last_updated": datetime.now().isoformat()
    }

def save_stats(stats):
    try:
        stats["last_updated"] = datetime.now().isoformat()
        with open(STATS_FILE, 'w') as f:
            json.dump(stats, f, indent=2)
    except Exception as e:
        log_event("STATS_SAVE_ERROR", error=str(e))

def update_stats(questions_count, processing_time, questions_data=None):
    stats = load_stats()
    stats["documents_processed"] += 1
    stats["questions_extracted"] += questions_count
    stats["total_processing_time"] += processing_time

    if questions_count > 0 and questions_data:
        total_conf = 0.0
        n = 0
        for q in questions_data:
            c = q.get('confidence')
            if isinstance(c, (int, float)):
                total_conf += float(c)
                n += 1
        stats["accuracy_rate"] = round((total_conf / n) * 100, 1) if n else stats.get("accuracy_rate", 85.0)
    else:
        stats["accuracy_rate"] = max(stats.get("accuracy_rate", 0.0), 85.0)

    save_stats(stats)
    return stats

# =========================================================
# Request logging
# =========================================================
@app.before_request
def _log_in():
    log_event("HTTP_IN", method=request.method, path=request.path)

@app.after_request
def _log_out(response):
    log_event("HTTP_OUT", status=response.status_code, path=request.path)
    return response

# =========================================================
# Routes
# =========================================================
@app.route('/')
def index():
    # Probe extractor readiness
    health = make_api_request('GET', 'health/ready')
    if health and health.status_code == 200:
        log_event("UPSTREAM_HEALTH_OK", endpoint="extract/health/ready")
    else:
        log_event("UPSTREAM_HEALTH_FAIL", status=(health.status_code if health else 'no-response'))

    # Probe KB readiness + discover route (both just for logging/visibility)
    kb_status, _kb_text = kb_health()
    kb_route = discover_kb_route()

    stats = load_stats()
    log_event("ROUTE_INDEX",
              kb_status=kb_status,
              kb_route=kb_route or "(unset)",
              kb_base=KK_KB_BASE or "(unset)")
    return render_template('index.html', api_config=api_config, stats=stats)

# ---- Upload & Extract ----
@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        log_event("UPLOAD_MISSING_FILE")
        return jsonify({'error': 'No file provided'}), 400

    f = request.files['file']
    if f.filename == '':
        log_event("UPLOAD_EMPTY_FILENAME")
        return jsonify({'error': 'No file selected'}), 400
    if not allowed_file(f.filename):
        log_event("UPLOAD_INVALID_TYPE", filename=f.filename)
        return jsonify({'error': 'Invalid file type'}), 400

    filename = secure_filename(f.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    f.save(filepath)
    log_event("UPLOAD_SAVED", filename=filename, path=filepath)

    use_llm = request.form.get('use_llm', 'false').lower() in ('true', '1', 'yes', 'on')
    use_sync = request.form.get('use_sync', 'false').lower() in ('true', '1', 'yes', 'on')
    mode = request.form.get('mode', 'balanced')

    try:
        with open(filepath, 'rb') as doc:
            files = {'file': doc}
            data = {'use_llm': 'true' if use_llm else 'false', 'mode': mode}
            endpoint = 'extract/sync' if use_sync else 'extract/'
            log_event("EXTRACT_CALL", endpoint=endpoint, mode=mode, use_llm=use_llm, use_sync=use_sync)
            resp = make_api_request('POST', endpoint, files=files, data=data)

        if not resp:
            log_event("EXTRACT_NO_RESPONSE")
            return jsonify({'error': 'API request failed (no response)'}), 502

        if resp.status_code == 200:
            payload = resp.json()
            if use_sync and isinstance(payload, dict) and 'questions' in payload:
                job_id = f"sync_{int(time.time())}"
                job_data[job_id] = {
                    'status': 'completed',
                    'questions': payload['questions'],
                    'timestamp': datetime.now().isoformat()
                }
                update_stats(len(payload['questions']), 5, payload['questions'])
                log_event("EXTRACT_SYNC_OK", job_id=job_id, q=len(payload['questions']))
                return jsonify({'success': True, 'job_id': job_id, 'questions': payload['questions']})
            else:
                job_id = payload.get('job_id')
                if job_id:
                    job_data[job_id] = {'status': 'processing', 'timestamp': datetime.now().isoformat()}
                    log_event("EXTRACT_ASYNC_OK", job_id=job_id)
                    return jsonify({'success': True, 'job_id': job_id, 'async': True})
                else:
                    log_event("EXTRACT_MISSING_JOB_ID")
                    return jsonify({'error': 'No job ID returned from API'}), 502
        else:
            try:
                msg = resp.json()
            except Exception:
                msg = resp.text
            log_event("EXTRACT_HTTP_ERROR", status=resp.status_code, message=str(msg))
            return jsonify({'error': f"HTTP {resp.status_code}: {msg}"}), resp.status_code

    except Exception as e:
        log_event("UPLOAD_EXCEPTION", error=str(e))
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
                log_event("UPLOAD_CLEANED", path=filepath)
        except Exception as e:
            log_event("UPLOAD_CLEANUP_ERROR", error=str(e))

# ---- Polling / Jobs ----
@app.route('/api/poll/<job_id>')
def poll_job(job_id):
    if job_id not in job_data:
        log_event("POLL_UNKNOWN_JOB", job_id=job_id)
        return jsonify({'error': 'Job not found'}), 404

    if job_data[job_id].get('status') == 'completed':
        log_event("POLL_ALREADY_DONE", job_id=job_id)
        return jsonify(job_data[job_id])

    resp = make_api_request('GET', f'jobs/{job_id}')
    if not resp:
        log_event("POLL_API_FAIL", job_id=job_id)
        return jsonify({'error': 'Failed to poll job status'}), 502

    if resp.status_code == 404:
        q = make_api_request('GET', f'jobs/{job_id}/questions')
        if q and q.status_code == 200:
            questions = q.json()
            job_data[job_id]['status'] = 'completed'
            job_data[job_id]['questions'] = questions
            update_stats(len(questions), 30, questions)
            log_event("POLL_RECOVERED_VIA_QUESTIONS", job_id=job_id, q=len(questions))
            return jsonify({'status': 'completed', 'questions': questions})
        log_event("POLL_NOT_FOUND_COMPLETING_EMPTY", job_id=job_id)
        return jsonify({'status': 'completed', 'questions': []})

    if resp.status_code != 200:
        try:
            msg = resp.json()
        except Exception:
            msg = resp.text
        log_event("POLL_HTTP_ERROR", job_id=job_id, status=resp.status_code, message=str(msg))
        return jsonify({'error': f"HTTP {resp.status_code}: {msg}"}), resp.status_code

    data = resp.json()
    job_data[job_id].update(data)

    status = data.get('status', 'processing')
    if status == 'completed':
        if data.get('questions') is not None:
            qs = data['questions']
            update_stats(len(qs), 30, qs)
            log_event("POLL_COMPLETED_INLINE", job_id=job_id, q=len(qs))
            return jsonify(data)
        q = make_api_request('GET', f'jobs/{job_id}/questions')
        if q and q.status_code == 200:
            qs = q.json()
            job_data[job_id]['questions'] = qs
            update_stats(len(qs), 30, qs)
            data['questions'] = qs
            log_event("POLL_COMPLETED_FETCHED", job_id=job_id, q=len(qs))
            return jsonify(data)
        log_event("POLL_COMPLETED_NO_QS", job_id=job_id)
        return jsonify({'status': 'completed', 'questions': []})

    if status == 'processing':
        start = job_data[job_id].get('timestamp')
        if start:
            try:
                started = datetime.fromisoformat(start)
                if datetime.now() - started > timedelta(minutes=10):
                    job_data[job_id]['status'] = 'timeout'
                    log_event("POLL_TIMEOUT", job_id=job_id)
                    return jsonify({'status': 'timeout', 'error': 'Processing timeout; try sync or no-LLM'})
            except Exception:
                pass
        return jsonify({'status': 'processing', 'progress': data.get('progress')})

    if status == 'failed':
        log_event("POLL_FAILED", job_id=job_id, data=data)
        return jsonify(data)

    return jsonify({'status': 'processing'})

@app.route('/api/job/<job_id>/questions')
def get_job_questions(job_id):
    if job_id not in job_data:
        return jsonify({'error': 'Job not found'}), 404

    job = job_data[job_id]
    if job.get('status') == 'completed' and 'questions' in job:
        return jsonify(job['questions'])

    resp = make_api_request('GET', f'jobs/{job_id}/questions')
    if resp and resp.status_code == 200:
        questions = resp.json()
        job['questions'] = questions
        return jsonify(questions)
    return jsonify({'error': 'Failed to fetch questions'}), 500

# ---- Review page ----
@app.route('/review/<job_id>')
def review(job_id):
    if job_id not in job_data:
        return redirect(url_for('index'))

    job = job_data[job_id]
    questions = job.get('questions') or []
    if not questions:
        resp = make_api_request('GET', f'jobs/{job_id}/questions')
        if resp and resp.status_code == 200:
            questions = resp.json()
            job['questions'] = questions

    return render_template('review.html', job_id=job_id, questions=questions, job=job)

@app.route('/api/save_edits/<job_id>', methods=['POST'])
def save_edits(job_id):
    data = request.get_json() or {}
    incoming = data.get('questions', [])
    if job_id not in job_edits:
        job_edits[job_id] = []

    index_map = {q.get('qid'): i for i, q in enumerate(job_edits[job_id])}
    for q in incoming:
        qid = q.get('qid')
        if qid in index_map:
            job_edits[job_id][index_map[qid]] = q
        else:
            job_edits[job_id].append(q)

    log_event("SAVE_EDITS", job_id=job_id, count=len(incoming))
    return jsonify({'success': True})

# ---- Export questions (docx/xlsx) ----
@app.route('/api/export/<job_id>/<format>')
def export_data(job_id, format):
    if job_id not in job_data:
        return jsonify({'error': 'Job not found'}), 404

    original = job_data[job_id].get('questions', []) or []
    edits = job_edits.get(job_id, []) or []
    edits_by_id = {q.get('qid'): q for q in edits}

    def normalize(q):
        q = dict(q or {})
        q['qid'] = q.get('qid') or q.get('id') or q.get('question_id')
        status = (q.get('status') or '').strip().lower()
        if status not in ('approved', 'pending', 'rejected'):
            status = 'pending'
        q['status'] = status
        try:
            c = float(q.get('confidence', 0))
            if c > 1:
                c = c / 100.0
        except Exception:
            c = 0.0
        q['confidence'] = round(c, 3)
        return q

    merged = []
    for oq in original:
        base = normalize(oq)
        override = normalize(edits_by_id.get(base.get('qid'), {}))
        for k in ('text', 'confidence', 'status', 'type', 'section_path', 'numbering', 'category'):
            v = override.get(k)
            if v not in (None, '', []):
                base[k] = v
        merged.append(base)

    approved_only = request.args.get('approved') in ('true', '1', 'yes', 'on')
    high_conf = request.args.get('high_confidence') in ('true', '1', 'yes', 'on')

    def keep(q):
        if approved_only and q.get('status') != 'approved':
            return False
        if high_conf and float(q.get('confidence', 0)) < 0.8:
            return False
        return True

    questions = [q for q in merged if keep(q)]

    if format == 'docx':
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('RFP Questions Report', 0)
            doc.add_paragraph(f'Generated on: {datetime.now():%Y-%m-%d %H:%M:%S}')
            doc.add_paragraph(f'Job ID: {job_id}')
            doc.add_paragraph(f'Total Questions: {len(questions)}')
            doc.add_paragraph('')

            for i, q in enumerate(questions, 1):
                doc.add_heading(f'Question {i}', level=2)
                doc.add_paragraph(f'Text: {q.get("text","N/A")}')
                doc.add_paragraph(f'Status: {q.get("status","pending").title()}')
                doc.add_paragraph(f'Confidence: {int(round(float(q.get("confidence",0))*100))}%')
                doc.add_paragraph(f'Type: {q.get("type","N/A")}')
                section = (q.get("section_path") or ["N/A"])[0] if q.get("section_path") else "N/A"
                doc.add_paragraph(f'Section: {section}')
                doc.add_paragraph('')

            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            fname = f"RFP_Questions_Report_{datetime.now():%Y%m%d_%H%M%S}.docx"
            return send_file(
                bio, as_attachment=True, download_name=fname,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except ImportError:
            content = [
                "RFP Questions Report",
                f"Generated on: {datetime.now():%Y-%m-%d %H:%M:%S}",
                f"Job ID: {job_id}",
                f"Total Questions: {len(questions)}",
                ""
            ]
            for i, q in enumerate(questions, 1):
                section = (q.get("section_path") or [""])[0] if q.get("section_path") else ""
                content.append(f"{i}. {q.get('text','N/A')}")
                content.append(f"   Status: {q.get('status','pending').title()}")
                content.append(f"   Confidence: {int(round(float(q.get('confidence',0))*100))}%")
                content.append(f"   Type: {q.get('type','N/A')}")
                content.append(f"   Section: {section}")
                content.append("")
            body = "\n".join(content)
            fname = f"RFP_Questions_Report_{datetime.now():%Y%m%d_%H%M%S}.txt"
            return body, 200, {
                'Content-Type': 'text/plain',
                'Content-Disposition': f'attachment; filename="{fname}"'
            }

    elif format == 'xlsx':
        try:
            import pandas as pd
            import openpyxl
            df = pd.DataFrame(questions)
            cols = ['qid','text','status','confidence','type','numbering','category','section_path']
            existing = [c for c in cols if c in df.columns]
            df = df[existing + [c for c in df.columns if c not in existing]]
            bio = io.BytesIO()
            with pd.ExcelWriter(bio, engine='openpyxl') as w:
                df.to_excel(w, sheet_name='RFP Questions', index=False)
            bio.seek(0)
            fname = f"RFP_Questions_Report_{datetime.now():%Y%m%d_%H%M%S}.xlsx"
            return send_file(
                bio, as_attachment=True, download_name=fname,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
        except ImportError:
            import csv
            out = io.StringIO()
            rows = []
            for q in questions:
                rows.append({
                    'qid': q.get('qid',''),
                    'text': q.get('text',''),
                    'status': q.get('status',''),
                    'confidence': q.get('confidence',''),
                    'type': q.get('type',''),
                    'section': (q.get('section_path') or [''])[0] if q.get('section_path') else '',
                    'numbering': q.get('numbering',''),
                    'category': q.get('category','')
                })
            if rows:
                writer = csv.DictWriter(out, fieldnames=rows[0].keys())
                writer.writeheader()
                writer.writerows(rows)
            fname = f"RFP_Questions_Report_{datetime.now():%Y%m%d_%H%M%S}.csv"
            return out.getvalue(), 200, {
                'Content-Type': 'text/csv',
                'Content-Disposition': f'attachment; filename="{fname}"'
            }

    return jsonify({'error': 'Invalid format'}), 400

# ===================== NEW: KB + RESPONSES =====================
@app.route('/api/approved_questions/<job_id>')
def get_approved_questions(job_id):
    if job_id not in job_data:
        return jsonify({'error':'Job not found'}), 404

    def merge_questions_with_edits(job_id_):
        original = job_data.get(job_id_, {}).get('questions', []) or []
        edits = job_edits.get(job_id_, []) or []
        edits_by_id = {q.get('qid') or q.get('id') or q.get('question_id'): q for q in edits}
        def norm(q):
            q = dict(q or {})
            q['qid'] = q.get('qid') or q.get('id') or q.get('question_id')
            status = (q.get('status') or '').strip().lower()
            if status not in ('approved', 'pending', 'rejected'):
                status = 'pending'
            q['status'] = status
            return q
        merged = []
        for oq in original:
            base = norm(oq)
            ov = edits_by_id.get(base['qid'])
            if ov:
                ov = norm(ov)
                for k in ('text','confidence','status','type','section_path','numbering','category'):
                    if ov.get(k) not in (None, '', []):
                        base[k] = ov.get(k)
            merged.append(base)
        return merged

    merged = get_approved_questions._merge_cache = locals().get('merge_questions_with_edits')(job_id)
    approved = [q for q in merged if q.get('status') == 'approved']
    return jsonify(approved)

# ---- KB proxy (DROP-IN) ----
@app.route('/api/kb/query-one', methods=['POST'])
def kb_query_one():
    if not KK_KB_BASE:
        log_event("KB_CALL_BLOCKED", reason="KB base not configured")
        return jsonify({'error': 'Knowledge Base not configured: set KK_KB_CONTENT_GUID or KK_KB_BASE'}), 503

    route = discover_kb_route()
    if not route:
        log_event("KB_CALL_BLOCKED", reason="KB route unresolved")
        return jsonify({'error': 'KB route could not be determined. Set KK_KB_ROUTE or ensure /openapi.json is accessible.'}), 404

    payload = request.get_json() or {}

    # Be generous with field names to match various KB schemas
    question = payload.get("query") or payload.get("question") or ""
    top_k = int(payload.get("max_results", payload.get("top_k", 5)))
    temperature = float(payload.get("temperature", 0.1))
    max_tokens = int(payload.get("max_tokens", 1500))
    conversation_history = payload.get("conversation_history") or []
    model = payload.get("model")  # optional passthrough

    body = {
        "query": question,
        "question": question,       # include both for compatibility
        "top_k": top_k,
        "max_results": top_k,       # include both for compatibility
        "temperature": temperature,
        "max_tokens": max_tokens,
        "conversation_history": conversation_history
    }
    if model:
        body["model"] = model

    url = f"{KK_KB_BASE}{route}"
    headers = get_kb_auth_header()

    log_event("KB_REQUEST_OUT",
              url=url, route=route, kb_base=KK_KB_BASE,
              hdr=list(headers.keys()), body_keys=list(body.keys()))

    try:
        r = requests.post(url, json=body, headers=headers, timeout=KK_KB_TIMEOUT)
    except requests.exceptions.RequestException as e:
        log_event("KB_REQUEST_ERR", error=str(e))
        return jsonify({'error': f'KB request failed: {e}'}), 502

    log_event("KB_REQUEST_BACK", status=r.status_code, ct=r.headers.get('Content-Type',''))

    if r.status_code == 404:
        return jsonify({'error': f'KB route not found on this content (404). Tried {route}. '
                                 f'Override KK_KB_ROUTE or check /openapi.json.'}), 404

    if r.status_code >= 400:
        try:
            return jsonify(r.json()), r.status_code
        except Exception:
            return jsonify({'error': r.text}), r.status_code

    try:
        return jsonify(r.json()), 200
    except Exception:
        return jsonify({"raw": r.text}), 200

# ---- Save edited KB responses ----
@app.route('/api/responses/save/<job_id>', methods=['POST'])
def save_responses(job_id):
    data = request.get_json() or {}
    incoming = data.get('responses', [])
    if job_id not in job_responses_edits:
        job_responses_edits[job_id] = []

    index = {r.get('qid'): i for i, r in enumerate(job_responses_edits[job_id])}
    for r in incoming:
        qid = r.get('qid')
        if qid in index:
            job_responses_edits[job_id][index[qid]] = r
        else:
            job_responses_edits[job_id].append(r)
    log_event("SAVE_RESPONSES", job_id=job_id, count=len(incoming))
    return jsonify({'success': True})

# ---- Responses UI page ----
@app.route('/responses/<job_id>')
def responses(job_id):
    if job_id not in job_data:
        return redirect(url_for('index'))
    if job_id not in job_responses:
        job_responses[job_id] = []
    return render_template('llm_responses.html', job_id=job_id)

# ---- Export KB responses (docx/xlsx) ----
@app.route('/api/export_responses/<job_id>/<format>')
def export_responses(job_id, format):
    if job_id not in job_data:
        return jsonify({'error': 'Job not found'}), 404

    # base from approved questions
    def merge_questions_with_edits(job_id_):
        original = job_data.get(job_id_, {}).get('questions', []) or []
        edits = job_edits.get(job_id_, []) or []
        edits_by_id = {q.get('qid') or q.get('id') or q.get('question_id'): q for q in edits}
        def norm(q):
            q = dict(q or {})
            q['qid'] = q.get('qid') or q.get('id') or q.get('question_id')
            status = (q.get('status') or '').strip().lower()
            if status not in ('approved', 'pending', 'rejected'):
                status = 'pending'
            q['status'] = status
            return q
        merged = []
        for oq in original:
            base = norm(oq)
            ov = edits_by_id.get(base['qid'])
            if ov:
                ov = norm(ov)
                for k in ('text','confidence','status','type','section_path','numbering','category'):
                    if ov.get(k) not in (None, '', []):
                        base[k] = ov.get(k)
            merged.append(base)
        return merged

    approved = [q for q in merge_questions_with_edits(job_id) if q.get('status') == 'approved']

    edits = job_responses_edits.get(job_id, []) or []
    by_id = {e.get('qid'): e for e in edits}

    rows = []
    for q in approved:
        qid = q.get('qid')
        e = by_id.get(qid, {})
        rows.append({
            'qid': qid,
            'question': q.get('text', ''),
            'response': e.get('response',''),
            'status': e.get('status','ok'),
            'type': q.get('type',''),
            'section': (q.get('section_path') or [''])[0] if q.get('section_path') else '',
            'sources': e.get('sources', [])
        })

    only_success = request.args.get('success') in ('true','1','yes','on')
    if only_success:
        rows = [r for r in rows if r.get('status') == 'ok']

    id_str = request.args.get('qids')
    if id_str:
        wanted = set(id_str.split(','))
        rows = [r for r in rows if r.get('qid') in wanted]

    if format == 'docx':
        try:
            from docx import Document
            from datetime import datetime as _dt
            doc = Document()
            doc.add_heading('RFP KB Responses', 0)
            doc.add_paragraph(f'Generated on: {_dt.now():%Y-%m-%d %H:%M:%S}')
            doc.add_paragraph(f'Job ID: {job_id}')
            doc.add_paragraph(f'Total Responses: {len(rows)}')
            doc.add_paragraph('')

            for i, r in enumerate(rows, 1):
                doc.add_heading(f'Item {i}', level=2)
                doc.add_paragraph(f'Question: {r['question']}")
                doc.add_paragraph(f"Response: {r.get('response') or '(empty)'}")
                doc.add_paragraph(f"Status: {r.get('status','ok').title()}")
                doc.add_paragraph(f"Type: {r.get('type') or 'N/A'}")
                doc.add_paragraph(f"Section: {r.get('section') or 'N/A'}")
                if r['sources']:
                    doc.add_paragraph('Sources:')
                    for s in r['sources']:
                        name = s.get('name') or s.get('uri') or 'source'
                        uri = s.get('uri')
                        line = f"- {name}" + (f" ({uri})" if uri else "")
                        doc.add_paragraph(line)
                doc.add_paragraph('')

            bio = io.BytesIO()
            doc.save(bio); bio.seek(0)
            fname = f"RFP_KB_Responses_{_dt.now():%Y%m%d_%H%M%S}.docx"
            return send_file(bio, as_attachment=True, download_name=fname,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        except ImportError:
            return jsonify({'error':'python-docx not installed on server'}), 500

    elif format == 'xlsx':
        try:
            import pandas as pd
            import openpyxl
            import io as _io, json as _json
            flat = []
            for r in rows:
                flat.append({
                    'qid': r['qid'],
                    'question': r['question'],
                    'response': r['response'],
                    'status': r['status'],
                    'type': r['type'],
                    'section': r['section'],
                    'sources_json': _json.dumps(r['sources'])
                })
            df = pd.DataFrame(flat)
            bio = _io.BytesIO()
            with pd.ExcelWriter(bio, engine='openpyxl') as w:
                df.to_excel(w, sheet_name='KB Responses', index=False)
            bio.seek(0)
            from datetime import datetime as _dt
            fname = f"RFP_KB_Responses_{_dt.now():%Y%m%d_%H%M%S}.xlsx"
            return send_file(bio, as_attachment=True, download_name=fname,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        except ImportError:
            return jsonify({'error':'pandas/openpyxl not installed on server'}), 500

    return jsonify({'error': 'Invalid format'}), 400

# ---- Quick KB config debug ----
@app.route('/api/kb/debug')
def kb_debug():
    route = discover_kb_route()
    return jsonify({
        "KK_KB_BASE": KK_KB_BASE or "(unset)",
        "KK_KB_CONTENT_GUID": KK_KB_CONTENT_GUID or "(unset)",
        "KK_KB_ROUTE": route or "(unset)",
        "KB_AUTH_HEADER_NAME": KK_KB_KEY_HEADER_NAME,
        "KB_AUTH_SCHEME": KK_KB_AUTH_SCHEME,
        "KB_KEY_PRESENT": bool(KK_KB_API_KEY),
    })

# =========================================================
# Entrypoint
# =========================================================
if __name__ == '__main__':
    print("Starting RFP Extraction App on port 5002...")
    print("Open your browser at http://localhost:5002")
    app.run(debug=True, port=5002, host='0.0.0.0')
